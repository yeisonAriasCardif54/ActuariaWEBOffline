from django.shortcuts import render
from django.http import JsonResponse
import pandas as pd
import traceback
import numpy as np
import docx
from docx.shared import Pt
import xlrd
import os
import time
from django.core.files.storage import FileSystemStorage
from pricing.views_.tarificador_libraries.register import create_register, update_register, get_categories, get_register_all, update_favorite, change_state


def index(request):
    username = ''
    email = ''
    first_name = ''
    last_name = ''
    if request.user.is_authenticated:
        username = request.user.username
        email = request.user.email
        first_name = request.user.first_name
        last_name = request.user.last_name

    configurationView = {
        'username': username,
        'email': email,
        'first_name': first_name,
        'last_name': last_name,
        'title': 'Cargar tarificador',
        'area': 'Pricing',
        'herramienta': 'uploadTarificador',
        'file': 'pricing/tarificador.html',
    }
    return render(request, "principal/base.html", configurationView)


def preUpload(request):
    file = ''
    insert_id = 0
    if request.method == 'POST' and request.FILES.get('file', False):
        myFile = request.FILES['file']
        filename = myFile.name
        file_ext = filename[-4:]
        if file_ext == 'xlsx':
            try:
                file, total_time = updateData(myFile)
                status = 1
                message = 'Tarificador cargado con éxito.'

                # -- Se guarda el archivo
                folder = 'static/pricing/tarificadores_inputs'
                fs = FileSystemStorage(location=folder)
                filename = fs.save(myFile.name, myFile)
                # file_url = fs.url(filename)
                # file = folder + '/' + filename
                registro, insert_id = create_register(request, filename)
                if registro < 1:
                    status = 0
                    message = 'ERROR:<br><br>Se presento un error al intentar guardar el intento (en BD).'
                    file = ''
                    error = traceback.format_exc()
                    update_register(insert_id=insert_id, file_output='Error al generar archivo', success=0, total_time=0, error=error)
                else:
                    update_register(insert_id=insert_id, file_output=file, success=1, total_time=total_time, error='')
            except:
                status = 0
                error = traceback.format_exc()
                message = 'ERROR:<br><br>' + error
        else:
            status = 0
            message = "ERROR:<br><br>Por favor seleccione el archivo en formato 'Libro de Excel (*.xlsx)'."
    else:
        status = 0
        message = 'ERROR:<br>Por favor seleccione un archivo.'
    data = {"status": status, "message": message, 'file': file, 'insert_id': insert_id}
    return JsonResponse(data, safe=False)


def upload(request):
    insert_id = 0
    file_return = ''
    folder = 'static/pricing/tarificadores_inputs'
    if request.method == 'POST' and request.FILES.get('file', False):
        myFilev1 = request.FILES['file']
        myFile = request.FILES['file']
        filename = myFile.name
        file_ext = filename[-4:]
        if file_ext == 'xlsx':
            registro, insert_id = create_register(request, filename)
            if registro < 1:
                status = 0
                message = 'ERROR:<br><br>Se presento un error al intentar guardar el intento (Log de eventos).'
                file_return = ''
            else:
                fs = FileSystemStorage(location=folder)
                filename = fs.save(myFile.name, myFile)
                file_url = fs.url(filename)
                file = folder + '/' + filename
                try:
                    file_return = updateData(myFilev1)
                    status = 1
                    message = 'Tarificador cargado con éxito.'
                except:
                    status = 0
                    error = traceback.format_exc()
                    message = 'ERROR:<br><br>' + error
        else:
            status = 0
            message = 'ERROR:<br><br>Por favor seleccione el archivo en formato Excel (.xlsx)'
            file_return = ''
    else:
        status = 0
        message = 'ERROR:<br><br>Por favor seleccione un archivo.'
        file_return = ''

    data = {"status": status, "message": message, "file": file_return, "insert_id": insert_id}
    return JsonResponse(data, safe=False)


def updateData(myFile):
    start_time_GLOBAL = time.time()

    # Input Excel
    workbook = xlrd.open_workbook(file_contents=myFile.read())

    # Worksheet 1
    worksheet = workbook.sheet_by_name('Tarificador')

    # Worksheet 2
    worksheet2 = workbook.sheet_by_name('Input')

    def prepare_worksheet(worksheet):

        num_cols = worksheet.ncols
        num_rows = worksheet.nrows - 1
        # Extract Rows Worksheet
        data = [worksheet.row_values(i) for i in range(0, num_rows)]

        # Create tempo Columns Worksheet
        columns = []
        months = range(0, num_cols)
        k = 1
        for month in months:
            colname = 'col%d' % month
            columns.append(colname)
            k += 1
        # Dataframe to work
        df = pd.DataFrame(data, columns=columns)
        # Remove empty rows
        df['e'] = df.apply(lambda x: "".join(x.astype(str)), axis=1)
        df = df[df['e'] != ''];
        df.drop('e', axis=1, inplace=True)
        del df['col0']
        return df

    # df_01
    df_01 = prepare_worksheet(worksheet)
    df_01['Marca'] = np.where(((df_01.col1 == 'Valor Asegurado') | (df_01.col1 == 'Severidad') |
                               (df_01.col1 == 'frecuencias') | (df_01.col1 == 'Incremento(+) / Descuento(-)') |
                               (df_01.col1 == 'Distribucion Prima Neta antes de IVA') | (df_01.col1 == 'IVA') |
                               (df_01.col1 == 'Prima bruta de IVA')), 1, 0)

    df_01['Marca2'] = df_01['Marca'].cumsum()
    df_01.drop('Marca', axis=1, inplace=True)
    # df_01.to_excel('C:/Users/b89591/Desktop/df_01.xlsx', index = None)

    # df_02
    df_02 = prepare_worksheet(worksheet2)

    # open an existing document
    path = os.path.join(os.path.dirname(os.path.dirname(__file__)), '../static/pricing/business_plan_templates/input.docx')
    doc = docx.Document(path)

    def create_table(df, df_inp, inf, sup):

        # print('inf',inf)

        # **************************************************************************
        # Data Valor Asegurado
        df_a = df[df.Marca2 == 1]
        # n drops the first n rows.
        df_a = df_a.iloc[1:]
        df_a.drop('Marca2', axis=1, inplace=True)
        # Trasponer Tabla
        df_a = df_a.set_index('col1').T

        # Data Severidad
        df_b = df[df.Marca2 == 2]
        # n drops the first n rows.
        df_b = df_b.iloc[1:]
        df_b.drop('Marca2', axis=1, inplace=True)
        # Trasponer Tabla
        df_b = df_b.set_index('col1').T

        # Data Prima Neta antes de IVA
        df_d = df[df.Marca2 == 5]
        # n drops the first n rows.
        df_d = df_d.iloc[1:]
        df_d.drop('Marca2', axis=1, inplace=True)
        # Trasponer Tabla
        df_d = df_d.set_index('col1').T

        # Data Prima bruta de IVA
        df_e = df[df.Marca2 == 7]
        # n drops the first n rows.
        df_e = df_e.iloc[1:]
        df_e.drop('Marca2', axis=1, inplace=True)
        # Trasponer Tabla
        df_e = df_e.set_index('col1').T

        # Data --- IVA ---
        df_f = df[df.Marca2 == 6]
        # n drops the first n rows.
        df_f = df_f.iloc[1:]
        # print(df_f)
        df_f.drop('Marca2', axis=1, inplace=True)
        # Trasponer Tabla
        df_f = df_f.set_index('col1').T
        # df_f.transform(lambda x: x + '1')

        # Data Incentivos
        df_inc = df[['col5', 'col6']]
        # n drops the first n rows.
        df_inc = df_inc.iloc[0:9]
        # print(df_inc)
        df_inc.rename(columns={'col5': 'Canal', 'col6': 'Valor'}, inplace=True)
        # df_inc.to_excel('C:/Users/b89591/Desktop/df_inc.xlsx', index = None)

        # Data IVA
        df_iva = round(((df.loc[(df['col1'] == 'IVA_'), 'col2'].astype(float)) + 1), 2)
        # print(df_iva)

        # Data Tipo de Venta
        df_tip_vent = (df_inp.loc[(df_inp['col1'] == 'Tipo de venta'), 'col2'])

        # Data Tipo de Intermediario
        df_tip_interm = (df_inp.loc[(df_inp['col1'] == 'Intermediario'), 'col2'])

        # **************************************************************************
        # Filter DataFrames
        df_a = df_a[inf:sup]
        df_b = df_b[inf:sup]
        df_d = df_d[inf:sup]
        df_e = df_e[inf:sup]
        df_f = df_f[inf:sup]

        # Columnas con Informacion
        dt2 = df_a.loc[:, df_a.ne(0).any()]
        dt2 = dt2.loc[:, df_a.ne('').any()]
        # dt2.to_excel('C:/Users/b89591/Desktop/dt2.xlsx', index = None)

        columnas_ = list(dt2)

        tempo1 = {};
        tempo2 = {};
        tempo3 = {}

        for columnas_a in columnas_:

            # Unir Valor Asegurado - Severidad
            if (columnas_a != 'Ecosistema'):
                df_a[columnas_a] = df_b[columnas_a].astype(str) + ' cuotas de ' + df_a[columnas_a].astype(str)
            else:
                df_a[columnas_a] = ''
            # print(df_a)

            # Filter Valores
            df_a = df_a.loc[:, df_a.ne(0).any()]
            df_a = df_a.loc[:, df_a.ne('').any()]

            # Data Prima Neta antes de IVA
            tempo1.update({k: v for k, v in zip([columnas_a], df_d[columnas_a])})

            # Data Prima bruta de IVA
            tempo2.update({k: v for k, v in zip([columnas_a], df_e[columnas_a])})

            # Data Aplica IVA
            tempo3.update({k: v for k, v in zip([columnas_a], df_f[columnas_a])})

            # Function Decrease 1 --- IVA ---

        def decrease_by_one(d):
            for key in d:
                d[key] = round((d[key] - 1), 2)
            return d

        decrease_by_one(tempo3)

        ##################################### TEMPORAL #####################################################
        dt_temp1 = dt2.copy()
        dt_temp1.loc[:] = ''
        dt_temp2 = dt_temp1.append(dt_temp1)
        dt_temp3 = dt_temp2.append(dt_temp1)
        dt_temp5 = dt_temp3.append(dt_temp2)
        dt_temp6 = dt_temp5.append(df_a)
        df_a = dt_temp6.append(dt_temp3)
        ##################################### TEMPORAL #####################################################

        # Fill Nan
        df_a = df_a.fillna('')
        # Concat Data
        df_ac = pd.concat([df_a, pd.DataFrame([tempo1]), pd.DataFrame([tempo2]), pd.DataFrame([tempo3])])
        # print('PARTE A')
        # df_ac.to_excel('C:/Users/b89591/Desktop/df_ac1.xlsx', index = None)

        # **************************************************************************
        # Data Edad mínima suscripción
        df_c = df[(df['col1'].isin(['Edad mínima suscripción', 'Edad permanencia']))]
        df_c.drop(['col1', 'Marca2'], axis=1, inplace=True)
        df_c = df_c.iloc[0:, inf]

        long_a = df_ac.shape[1]
        keys = list(df_ac.columns)

        valores = list(df_c)

        values = []
        for _ in range(long_a):
            values.append(valores)

        # values = [ item for item in valores for _ in range(long_a) ]
        new_dict = {k: v for k, v in zip(keys, values)}
        unir = pd.DataFrame(new_dict, columns=keys)
        unir['Ecosistema'] = ''

        frames = [df_ac, unir]
        df_ac = pd.concat(frames)
        # print('PARTE B')
        # df_ac.to_excel('C:/Users/b89591/Desktop/df_ac2.xlsx', index = None)

        # Definir Index Final
        index_ac = ['Ramo', 'Código NT', 'Carencia', 'Deducible', 'Exclusiones', 'Límite del beneficio', 'Número de Eventos', 'Tipo de Beneficio', 'Beneficiario',
                    'Peso de Cobertura antes de IVA', 'Peso de Cobertura después de IVA', 'Aplica IVA', 'Edades Limites (suscripción)', 'Edad Limites (beneficio)']
        # index_ac = ['Límite del beneficio','Peso de Cobertura antes de IVA','Peso de Cobertura después de IVA','Aplica IVA','Edades Limites (suscripción)','Edad Limites (beneficio)']
        df_ac.index = index_ac
        df_ac = df_ac.rename_axis('Coberturas').reset_index()

        # Move Ecosistema column to the end
        cols_shift = list(df_ac.columns)
        cols_shift = cols_shift[0:1] + cols_shift[2:] + cols_shift[1:2]  # + [cols_shift[-1]]
        df_ac = df_ac[cols_shift]

        # ***********************************************************************************************************************
        # Data Texto
        # ***********************************************************************************************************************

        df_x = df[(df['col1'].isin(['Código Producto', 'Linea', 'Canal', 'Prima Neta de IVA',
                                    'Prima Cliente', 'Client Value', 'Utilidad Técnica', 'Ecosistema antes de IVA',
                                    'Costo recaudo IVA incluido', 'Porcentaje de Gastos', 'Gastos', 'Comisión Banco', 'Comisión Intermediario', 'IVA_', 'Periodicidad pago prima']))]

        # df_x.to_excel('C:/Users/b89591/Desktop/df_x1.xlsx', index = None)
        # Columna Variable
        df_x_pre = df_x[['col1']].reset_index(drop=True)
        # print('df_x_pre',df_x_pre)

        # Valores
        df_x.drop(['col1', 'Marca2'], axis=1, inplace=True)  # ;df_x.to_excel('C:/Users/b89591/Desktop/df_x0.xlsx', index = None)

        df_x = df_x.iloc[0:, inf].reset_index(drop=True)

        # print('df_x',df_x)

        df_x = pd.concat([df_x_pre, df_x], axis=1)
        # print('df_x2',df_x)

        df_x = df_x.rename(columns={df_x.columns[1]: "value"})  # ;df_x.to_excel('C:/Users/b89591/Desktop/df_x1.xlsx', index = None)
        # print('PARTE C')

        producto = df_x.loc[(df_x['col1'] == 'Código Producto'), 'value']
        linea = df_x.loc[(df_x['col1'] == 'Linea'), 'value']
        canal = df_x.loc[(df_x['col1'] == 'Canal'), 'value']
        prima_ant = df_x.loc[(df_x['col1'] == 'Prima Neta de IVA'), 'value']
        prima_des = df_x.loc[(df_x['col1'] == 'Prima Cliente'), 'value']

        cli_val = df_x.loc[(df_x['col1'] == 'Client Value'), 'value']
        uti_tec = df_x.loc[(df_x['col1'] == 'Utilidad Técnica'), 'value']
        eco_iva = df_x.loc[(df_x['col1'] == 'Ecosistema antes de IVA'), 'value']
        costo_eco = df_x.loc[(df_x['col1'] == 'Costo recaudo IVA incluido'), 'value']

        comis_bk_net = df_x.loc[(df_x['col1'] == 'Comisión Banco'), 'value']
        # print('comis_bk_net',comis_bk_net)

        # print('df_iva.iloc[0]',df_iva.iloc[0])

        # print('df_iva.iloc[0]',df_iva.iloc[0])
        comis_bk_bru = str((round((((float(comis_bk_net.iloc[0])) * (float(df_iva.iloc[0]))) * 100), 2))) + ' %'
        # comis_bk_bru = str( (round ( (( (comis_bk_net.iloc[0]) * (df_iva.iloc[0]) ) *100) ,2 ) ) ) + ' %'

        comis_int_net = df_x.loc[(df_x['col1'] == 'Comisión Intermediario'), 'value']

        gastos = df_x.loc[(df_x['col1'] == 'Porcentaje de Gastos') + (df_x['col1'] == 'Gastos'), 'value']
        # gastos_ =   (str(  round(float (gastos.iloc[1]) ,0) * 100) + '%') + ' (' + (str(round(int(gastos.iloc[0]),0))) + ' COP)'
        gastos_ = (str(round((float(gastos.iloc[1]) * 100), 2)) + '%') + ' (' + (str(round(int(gastos.iloc[0]), 0))) + ' COP)'

        incentivo = df_inc.loc[df_inc['Canal'] == str(canal.iloc[0]), ['Canal', 'Valor']]
        incentivo_ = df_inc.loc[df_inc['Canal'] == str(canal.iloc[0]), ['Valor']]
        incentivo_ = incentivo_.values
        var_inc = np.where(incentivo.Canal != 'PDV',
                           'Costo de venta por Telemarketing (neto de IVA) : ' + str(float(incentivo_)),
                           'Incentivos (% de la prima mensual neta de IVA) : ' + (str((float(incentivo_) * 100)) + '%'),
                           )

        tipo_pago_prima = df_x.loc[(df_x['col1'] == 'Periodicidad pago prima'), 'value']

        if (str(df_tip_interm.iloc[0]) == 'Aseguradora'):

            com_alfa_net = (df_x.loc[(df_x['col1'] == 'Comisión Intermediario'), 'value'])
            com_alfa_net = (str(round((float(com_alfa_net.iloc[0]) * 100), 2)) + '%')

            com_alfa_bru = (df_x.loc[(df_x['col1'] == 'Comisión Intermediario'), 'value'])
            com_alfa_bru = (str(round((float(com_alfa_bru.iloc[0]) * 100), 2)) + '%')

            com_interm_bru = str(0)
            com_interm_net = str(0)
        else:
            com_alfa_net = str(0)
            com_alfa_bru = str(0)

            com_interm_net = (df_x.loc[(df_x['col1'] == 'Comisión Intermediario'), 'value'])
            com_interm_net = (str(round((float(com_interm_net.iloc[0]) * 100), 2)) + '%')

            com_interm_bru = (df_x.loc[(df_x['col1'] == 'Comisión Intermediario'), 'value'])
            com_interm_bru = (str(round(((float(com_interm_bru.iloc[0]) * (float(df_iva.iloc[0]))) * 100), 2)) + '%')

        # **************************************************** DATA FINAL WORD *********************************************************
        # add Text

        paragraph_format = doc.styles['Normal'].paragraph_format
        paragraph_format.space_before

        paragraph_format.space_before = Pt(2)
        paragraph_format.space_before.pt

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        p = doc.add_paragraph()
        run = p.add_run('Producto: ' + str(producto.iloc[0]))
        # run.font.size = Pt(16)
        run.bold = True

        p = doc.add_paragraph('Modo de Venta: ' + str(df_tip_vent.iloc[0]))
        # p_format.line_spacing = Pt(35)
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Línea de Crédito: ' + str(linea.iloc[0]))
        run = p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Canal de Venta: ' + str(canal.iloc[0]))
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Comisión del socio (neta de IVA): ' + (str(round(float(comis_bk_net.iloc[0]), 3) * 100) + '%'))
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Comisión del socio (bruto de IVA): ' + comis_bk_bru)
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Comisión ALFA (neta de IVA): ' + com_alfa_net)
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Comisión ALFA (bruto de IVA): ' + com_alfa_bru)
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Comisión Intermediario (neta de IVA): ' + com_interm_net)
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Comisión Intermediario (bruto de IVA): ' + com_interm_bru)
        p.add_run()  # ;    run.add_break()

        # Incentivo
        p = doc.add_paragraph(var_inc)
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Costo de ecosistema mensual (neta de IVA): ' + str(eco_iva.iloc[0]))
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Costo de recaudo (% prima neta): ' + str(costo_eco.iloc[0]))
        p.add_run()  # ;    run.add_break()

        #p = doc.add_paragraph('Gastos: ' + gastos_)
        #p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Frecuencia de pago de prima: ' + str(tipo_pago_prima.iloc[0]))
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Prima antes de IVA: ' + str(round((prima_ant.iloc[0]), 0)))
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Prima después de IVA: ' + str(round((prima_des.iloc[0]), 0)))
        p.add_run()  # ;    run.add_break()

        p = doc.add_paragraph('Client Value: ' + (str(round((float(cli_val.iloc[0]) * 100), 0)) + '%'))
        p.add_run()  # ;    run.add_break()

        #p = doc.add_paragraph('Utilidad Técnica: ' + (str(round((float(uti_tec.iloc[0]) * 100), 2)) + '%'))
        #run = p.add_run();
        run.add_break()

        # add a table to the end and create a reference variable
        # extra row is so we can add the header row
        t = doc.add_table(df_ac.shape[0] + 1, df_ac.shape[1])
        t.style = 'Table Grid'

        # add the header rows.
        for j in range(df_ac.shape[-1]):
            t.cell(0, j).text = df_ac.columns[j]
            # add the rest of the data frame
        for i in range(df_ac.shape[0]):
            for j in range(df_ac.shape[-1]):
                t.cell(i + 1, j).text = str(df_ac.values[i, j])
        doc.add_page_break()

        # print('PARTE D')
        # # print('----------------')

    df_long = df_01.iloc[9:10]
    df_long = df_long.drop(['Marca2'], axis=1)
    df_long = df_long.set_index('col1').T
    seriesObj = df_long.apply(lambda x: True if x['Código Producto'] != '' else False, axis=1)
    # Count number of True in series
    long = len(seriesObj[seriesObj == True].index)

    ini = 0
    xx = 1
    while ini < long:
        # # print(xx)
        create_table(df_01, df_02, ini, (ini + 1))
        ini += 1
        xx += 1

    # save the doc
    newPath = os.path.join(os.path.dirname(os.path.dirname(__file__)), '../static/pricing/tan_output/')
    nameFile = 'output_' + time.strftime("%Y_%m_%d_%H_%M_%S") + '.docx'
    finalFile = newPath + nameFile
    doc.save(finalFile)

    # -- Calcular tiempo total en generar TAN
    total_time = (time.time() - start_time_GLOBAL)
    total_time = format(total_time, '.2f')

    return nameFile, str(total_time)
